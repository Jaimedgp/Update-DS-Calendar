#
# ReadDocxFile 
#-----------------
# This script is a part of the Master-Schedule program to update the 
# master calendar automatically from the docx file to Google Drive
#
# This class is developed in order to read an expecific .docx file
# structured as a calendar with a table for each month of classes.
#
# author: Jaime Diez Gonzalez-Pardo (Jaimedgp)
# github: https://github.com/Jaimedgp
################################################################################

from docx import Document
from datetime import date

from classEvent import ClassEvent

class ReadDocxFile(object):

    """
        Read the information from the 'Master en Ciencia de Datos' calendar
        obtaining the classes information from the cell corresponding to a
        centain date
    """

    def __init__(self, doc_name):
        """
            Read a .docx file formed by tables (per month) as a calendar

            :doc_name: name of the .docx file 'path/to/file.docx'
            :date: date to read <class type='datetime'>
        """
        self._doc = Document(doc_name)


    def get_date_coordinates(self):
        """
            get atributees coordinates (num_month, num_week, day_week) for the
            date's cell
        """

        day_week  = self._date.weekday() + 1
        if day_week == 7: day_week = 0

        start_date = self._date.replace(day=1).isocalendar()
        if start_date[2] == 7:
            start_date = self._date.replace(day=2).isocalendar()

        num_week = (self._date.isocalendar()[1]
                        - start_date[1] + 2)

        num_month = self._date.month - 10

        if num_month < 0:
            num_month += 12

        return num_month, num_week, day_week


    def read_cell(self, date='today'):
        """
            Read the string of the coordinates' cell
        """

        if date == 'today':
            self._date = date.today()
        else:
            self._date = date

        coordinates = self.get_date_coordinates()

        cell_string =  (self._doc.tables[coordinates[0]].
                                         rows[coordinates[1]].
                                        cells[coordinates[2]].paragraphs
                       )

        dy_schedule = self.split_classes(
                                    [line.text for line in cell_string
                                                                if line.text != '']
                                   )

        classes = []

        for clss in dy_schedule:
            classes.append(ClassEvent(clss, self._date))

        return classes


    def split_classes(self, cell_info):
        """
            Some cells have two classes information separated by
            a '----' line, so it is necessary to split both strings
        """

        for i, item in enumerate(cell_info):
            if "----" in item:
                dy_schedule = [cell_info[1:i], cell_info[i+1:]]
                break
        else:
            dy_schedule = [cell_info[1:]]

        return dy_schedule


